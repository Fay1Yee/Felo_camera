#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os

try:
    from docx import Document
except ImportError:
    print("需要安装python-docx库: pip install python-docx")
    sys.exit(1)

def read_docx_file(file_path):
    """读取docx文件内容"""
    try:
        doc = Document(file_path)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        # 也读取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(" | ".join(row_text))
        
        return "\n".join(content)
    
    except Exception as e:
        print(f"读取文档时出错: {e}")
        return None

if __name__ == "__main__":
    file_path = "timestampcategoryconfidencereasons2025.docx"
    
    if not os.path.exists(file_path):
        print(f"文件不存在: {file_path}")
        sys.exit(1)
    
    content = read_docx_file(file_path)
    if content:
        print("=== 文档内容 ===")
        print(content)
        
        # 保存为文本文件以便后续处理
        with open("document_content.txt", "w", encoding="utf-8") as f:
            f.write(content)
        print("\n=== 内容已保存到 document_content.txt ===")
    else:
        print("无法读取文档内容")